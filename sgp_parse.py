#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# Parse the Singpho Dictionary to get .tsv file

from sgp_converter import sgp_converter

from good_results_sgp import good_results_sgp
from preconverted_assamese import sgp_assamese_info

from docx import Document

import os
import sys

parts_of_speech = ['v', 'n', 'adj', 'prep', 'pron', 'adv', 'inter', 'Inter', 'adj/adv', 'Exc', 'int', 'pl']

tnr = 'Times New Roman'
nirmala = 'Nirmala UI'
tanmatra = 'Tanmatra Boishakhi'

converter = sgp_converter()

from preconverted_assamese import sgp_assamese_info


def assamese_convert_dictionary():
    # Get preconverted as dictionary of input to index
    preconv_to_unicode = {}
    num_dups = 0
    for precon in sgp_assamese_info:
        str_index = '%d' % precon[0]
        precon_data = precon[2]
        try:
            unicode_out = good_results_sgp[str_index]
            if precon[2] in preconv_to_unicode:
                num_dups += 1
            preconv_to_unicode[precon[2]] = unicode_out
        except:
            preconv_to_unicode[precon[2]] = None

        # print('%s : %s' % (precon[0], preconv_to_unicode[precon_data]))
    # print('%d duplicates found' % num_dups)
    return preconv_to_unicode

# Global
conversion_dict = assamese_convert_dictionary()


def new_entry():
    return {'heading': [],
             'sgp_word': '',
             'definition': [],
             'def_list': [],  # a list of [part_of_speech, definition]
             'part_of_speech': [],
             'assamese1': [],
             'assamese2': []
             }

def process_assamese1(r_text, entry, entries, previous):
    entry[previous].append(r_text)
    return entry, previous


def process_assamese2(r_text, entry, entries, previous):
    entry[previous].append(r_text)
    return entry, previous


def process_definition(r_text, entry, entries, previous):
    # Looks like a definition
    if previous == 'part_of_speech' and r_text.strip() == '.':
        return entry, previous
    if len(r_text) == 1:
        # Patch single characters to previous entry
        if entry['definition']:
            entry['definition'][-1] += r_text
        else:
            entry['definition'].append(r_text)
    else:
        entry['definition'].append(r_text)
    # Update the latest
    if entry['def_list']:
        latest_pos = entry['def_list'][-1][0]
        last_def = entry['def_list'][-1][1]
        if last_def:
            last_def.append(r_text)
            entry['def_list'][-1][1] = last_def
        else:
            entry['def_list'][-1] = [latest_pos, [r_text]]
    else:
        # There should be something here!
        pass
    previous == 'definition'
    return entry, previous


def process_sgp_word(r_text, entry, entries, previous):
    if r_text[0].strip() == '(':
        entries.append(entry)
        entry = new_entry()
        alpha_group = r_text
        entry['heading'] = [alpha_group]
        print('ALPHABETIC GROUP: %s' % alpha_group)
        entries.append(entry)
        entry = new_entry()
    else:
        if entry:
            # New item:
            entries.append(entry)
        entry = new_entry()
        entry['sgp_word'] = r_text
        previous = 'sgp_word'
    return entry, previous


def process_part_of_speech(r_text, entry, entries, previous):
    # Part of speech or definition?
    entry['part_of_speech'].append(r_text + '.')
    previous = 'part_of_speech'
    entry['def_list'].append([r_text + '.', []])
    return entry, previous


def process_paragraphs(paragraphs):
    debug = False
    section_separator = '<><><>'
    index = 0
    main_dictionary = True
    entries = []
    entry = new_entry()
    # Set up headings
    entry['heading'] = ['Heading']
    entry['sgp_word'] = 'Word'
    entry['definition'] = ['Definitions']
    entry['part_of_speech'] = ["Part of Speech"]
    entry['assamese1'] = ["Assamese1"]
    entry['assamese2'] = ["Assamese2"]
    entries.append(entry)
    entry = new_entry()

    sections = []
    alpha_group = None
    headings = []
    for p in paragraphs:
        if debug:
            print('P %d: %d runs: %s' %(index, len(p.runs), p.text))
        run_index = 0
        p_text = p.text
        if p.text.find(section_separator) >= 0:
            # Output the current one.
            entries.append(entry)
            entry = new_entry()
            #entry['heading'] = 'Section %s' % (p.text)
            #entries.append(entry)
            print('SECTION: %s' % p.text)
            sections.append([p.text, index])
            entry = new_entry()
            continue

        if p.text.find('Words ') == 0:
            entries.append(entry)
            entry = new_entry()
            entry['heading'] = ['Info: %s' % (p.text)]
            headings.append([p.text, index])
            entries.append(entry)
            entry = new_entry()

        # Get groups of run data consolidated by style and font name
        combined = combine_runs(p)

        previous = None
        in_assamese_section = False
        for run_index in range(len(p.runs)):
            r = p.runs[run_index]
            r_text = r.text.strip()

            # Skip space-only
            if not r_text.replace(' ', '').replace('\t', ''):
                continue

            if run_index == 0 and r.font.bold and r.font.name == tnr:
                # New sgp entry
                entry, previous = process_sgp_word(r_text, entry, entries, previous)
                continue

            if (not r.font.italic and not r.font.bold and r.font.name == tnr and
                not in_assamese_section):
                entry, previous = process_definition(r_text, entry, entries, previous)
                continue

            if r.font.italic and r.font.name == tnr and r_text in parts_of_speech:
                # Part of speech or definition?
                entry, previous = process_part_of_speech(r_text, entry, entries, previous)
                continue

            if r.font.italic and r.font.name == tnr and previous == 'definition':
                # Add to the current definition
                if entry['def_list']:
                    latest_pos = entry['def_list'][-1][0]
                    last_def = entry['def_list'][-1][1] + r_text
                    entry['def_list'][-1] = [latest_pos, ['def_list'][-1][1].append(r_text)]
                else:
                    # ???
                    pass
                continue

            if r.font.name == tnr and in_assamese_section:
                # Add part of speech and other things while in an Assamese section
                entry, previous = process_assamese1(r_text, entry, entries, previous)
                continue

            if r.font.name == nirmala:
                previous = 'assamese1'
                entry, previous = process_assamese1(r_text, entry, entries, previous)
                in_assamese_section = True
                continue

            if r.font.name == tanmatra:
                previous = 'assamese2'
                in_assamese_section = True
                # TODO: CONVERT!!
                entry, previous = process_assamese1(r_text, entry, entries, previous)
                continue
            # What's this?
            pass
        index += 1


    return entries, sections, headings


def get_file(input_file):
    try:
        doc = Document(input_file)
    except BaseException as err:
        return None
    return doc

def combine_runs(p):
    # Combine runs by font and style
    result = []
    runs = p.runs
    p_text = p.text
    if not runs:
        return result
    current_style = None
    current_font_name = None
    current_run = []
    for r in runs:
        style = None
        if r.font.bold:
            style = 'bold'
        elif r.font.italic:
            style = 'ital'
        if style == current_style and r.font.name == current_font_name:
            current_run.append(r.text)
        else:
            # It changed
            all_text = ''.join(current_run)
            if True:
                result.append({'style': current_style,
                               'font_name': current_font_name,
                               'text': all_text
                               })
            current_run = [r.text]
            current_style = style
            current_font_name = r.font.name
    # The last bit
    all_text = ''.join(current_run)
    if True:
        result.append({'style': current_style,
                       'font_name': current_font_name,
                       'text': all_text
                       })
    return result



def save_entries(entries, folder):
    fname = 'SingPho_entries.tsv'
    path = os.path.join(folder, fname)
    count = 0
    unconverted_set = set()

    preconverted_assames_encoding = []
    with open(path, 'w') as out:
        for entry in entries:
            def_list = entry['def_list']
            def_strings = []
            for d in def_list:
                try:
                    def_joined = d[0] + ' '.join(d[1])
                    def_strings.append(def_joined)
                except BaseException as e:
                    pass
            definitions_out = '; '.join(def_strings)
            converted_assamese = ''
            if not entry['heading']:
                assamese2_encoded = ''.join(entry['assamese2'])
                preconverted_assames_encoding.append([entry['sgp_word'], assamese2_encoded])
                converted_assamese, unconverted = convert_assamese2(assamese2_encoded)

                # Check against good results
                if assamese2_encoded in conversion_dict:
                    expected = conversion_dict[assamese2_encoded]
                    if converted_assamese != expected:
                        print('Conversion not correct for %d: \"%s" != "%s"' % (count, converted_assamese, expected))
                        converted_assamese = '* ' + converted_assamese + ' *'
                else:
                    expected = None
                    converted_assamese = 'NO EXPECTED VALUE'

                # TODO: flag unconverted things
                if unconverted:
                    for key in unconverted.keys():
                        unconverted_set.add(key)
                    # print('!!! Unconverted: %s in entry %s' % (unconverted, entry))
            items = [
                ', '.join(entry['heading']),
                entry['sgp_word'],
                ', '.join(entry['part_of_speech']),
                definitions_out,
                ', '.join(entry['assamese1']),
                converted_assamese,
            ]
            try:
                out_line = '%s\n' % '\t'.join(items)
            except TypeError as e:
                pass
            out.write(out_line)
            count += 1

        print('Unconverted values: %s' % unconverted_set)
        char_list = []
        for item in unconverted_set:
            values = item.split(' ')
            char_list.append(values[1])
            print(' %s %s' % (values[1], values[2]))
        print('Unconverted chars = %s' % ('  '.join(char_list)))

        # Get the context for each converted value.
        path = os.path.join(folder, 'preconverted_assamese.txt')
        index = 0
        with open(path, 'w') as out:
            for text in preconverted_assames_encoding:
                if not text[1].strip():
                    # Nothing there
                    continue
                line_out = '[%4d, "%s", "%s"],\n' % (index, text[0], text[1])
                out.write(line_out)
                index += 1

        non_converted_keys = converter.not_converted_context.keys()
        for key in non_converted_keys:
            print('    [%s: %s]' % (key, converter.not_converted_context[key]))
        return count


def convert_assamese2(text):
    result = None
    if isinstance(text, list):
        result = []
        unconverted = None
        for item in text:
            new_text, unconverted = converter.convert(item, tanmatra)
            result.append(new_text)
    else:
        result, unconverted = converter.convert(text, tanmatra)
    return result, unconverted


# If bold and Times font -> may be start of a definition
# If italic and Times font -> may be part of speech or a definition
# if font is "Nirmala UI", its a Assamese definition
# if font is "Tanmatra Boishakhi", needs to be converted to Assamese

# A Run with only "." --> merge with previous run
# A Run with only spaces" -->  ?ignore?
#
def main(argv):
    print('Processing file with Tanmatra Boishakhi font!')

    input_file = argv[1]
    doc = get_file(input_file)
    print('doc %s' % doc)

    paragraphs = doc.paragraphs
    print('doc %d' % len(paragraphs))

    entries, headings, sections = process_paragraphs(doc.paragraphs)

    input_folder = os.path.dirname(input_file)
    num_entries = save_entries(entries, input_folder)
    return

if __name__ == "__main__":
    main(sys.argv)




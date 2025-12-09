#!/usr/bin/python3
# -*- coding: utf-8 -*-

# [START gae_python312_app]
from flask import Flask, render_template, stream_with_context, request, Response, send_file

# https://flask.palletsprojects.com/en/2.1.x/patterns/fileuploads/

import datetime

from io import BytesIO
from io import StringIO

import json
import logging

# https://github.com/saffsd/langid.py
import langid  # For identifying language of text

import random
from zipfile import ZipFile

import queue
import random
import threading
import time

import google.auth

from google.cloud import tasks

import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


import adlamConversion
import ahomConversion
import phkConversion
import mendeConverter

from convertDoc2 import ConvertDocx

import convertXls

# Global logger
logger = logging.getLogger('uploader')
logger.setLevel(logging.DEBUG)


# Dictionary of the converters by language code
converters = {}
converters['ff'] = adlamConversion.AdlamConverter()
converters['aho'] = ahomConversion.AhomConverter()
converters['phk'] = phkConversion.PhakeConverter()
converters['men'] = mendeConverter.MendeConverter()

lang_names_from_codes = {
    'ff': 'Poular',
    'aho': 'Tai Ahom',
    'phk': 'Tai Phake',
    'men': 'Mende Kikakui',
    'und': 'Unknown'
    }

# Datastore
#datastore_client = datastore.Client()

# If `entrypoint` is not defined in app.yaml, App Engine will look for an app
# called `app` in `main.py`.
app = Flask(__name__)

ts_client = tasks.CloudTasksClient()

_, PROJECT_ID = google.auth.default()
QUEUE_NAME = 'font-convert-queue'
REGION_ID = LOCATION_ID = 'us-central1'
QUEUE_PATH = ts_client.queue_path(PROJECT_ID, REGION_ID, QUEUE_NAME)

# Try using exporting threads to give progress report
exporting_threads = {}
      
# Global queue for messages
queue = queue.Queue(maxsize=100)

app.debug = True


@app.route('/')
def hello():
    """Top of the conversion application."""
    who = request.url
    return render_template('main.html', base=who)


# https://pythonbasics.org/flask-upload-files
@app.route('/upload/adlam')
def upload():
   who = request.host_url
   scriptIndex = request.args.get('scriptIndex', 0)
   # For indexing a thread
   taskId = random.randint(0, 7777)
   return render_template('upload.html',
                          base=who,
                          lang='ff',
                          scriptIndex=scriptIndex,
                          taskId=taskId
   )


@app.route('/upload_xlsx/adlam')
def upload_xslx():
   who = request.host_url
   scriptIndex = request.args.get('scriptIndex', 0)
   # For indexing a thread
   taskId = random.randint(0, 7777)
   return render_template('upload_xlsx.html',
                          base=who,
                          lang='ff',
                          scriptIndex=scriptIndex,
                          taskId=taskId
   )

# load file with explicit language and encoding
@app.route('/uploadlang')
def uploadLang():
    who = request.host_url
    lang = request.args.get('lang', 'und')
    script_index = request.args.get('script_index', 0)
    taskId = random.randint(0, 7777)
    try:
        lang_name = lang_names_from_codes[lang]
    except:
        lang_name = '??'
        
    unicode_font_list = ['Noto Sans', 'Noto Serif']
    if lang == 'aho':
        unicode_font_list = ['Noto Serif Ahom',
                             'Ahom Manuscript Unicode']
    elif lang == 'phk':
        unicode_font_list = ['Ramayana Unicode',
                             'Myanmar Text',
                             'Noto Sans Myanmar Regular',
                             'Noto Serif Myanmar Regular',
                             ]

    return render_template('upload_lang.html',
                           base=who,
                           lang=lang,
                           lang_name=lang_name,
                           taskId=taskId,
                           script_index=script_index,
                           fonts=unicode_font_list
    )


# Global
msgToSend = 'First Message'
countSent = 0

def read_file_chunks(fd):
  chunks = 0
  while 1:
      buf = fd.read(8192)
      if buf:
          yield buf
      else:
          break
      chunks += 1
  if app.debug:
      progressFn('Download complete')

# A way to create progress functions with other information neede
# for communication
class ProgressClass():
    def __init__(self, converter, thread=None):
        self.converter = converter
        self.thread = thread  # May be available
        self.status = "Nothing"

    def send(self, message):
        global queue
        # create output
        self.status = message
        self.thread.setStatus(message)
        queue.put(message)

    def stop_updates(self, final_msg):
        self.status = final_msg
        self.thread.setStatus(final_msg)
        queue.put(final_msg)
        queue.put('## STOP ##')
        queue.task_done()
        return
        

# Simple output function for tracking processing
def progressFn(msg):
    global msgToSend
    global countSent
    if app.debug:
        print('PROGRESS: %s' % msg)
    msgToSend = msg
    countSent = 1
    
def summarizeDoc(doc):
    # Find the fonts and languages of the paragraphs

    return

def findDocFonts(doc):
    fontsFound = {}
    lang_codes = {}
    if not doc:
        return fontsFound
    lang_paragraphs = {}
    if doc.paragraphs:
        fontsFound = getFontsInParagraphs(doc.paragraphs, fontsFound)
        lang_paragraphs = getLangsInParagraphs(doc.paragraphs, lang_codes, lang_paragraphs)

    for table in doc.tables:
        rows = table.rows
        for row in rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                fontsFount = getFontsInParagraphs(paragraphs, fontsFound)
                getLangsInParagraphs(paragraphs, lang_codes, lang_paragraphs)

    sections = doc.sections
    for section in sections:
        try:
            header = section.header
            fontsFound = getFontsInParagraphs(header.paragraphs, fontsFound)
            getLangsInParagraphs(header.paragraphs, lang_codes, lang_paragraphs)
        except:
            pass
        try:
            footer = section.footer
            fontsFound = getFontsInParagraphs(footer.paragraphs, fontsFound)
            getLangsInParagraphs(footer.paragraphs, lang_codes, lang_paragraphs)
        except:
            pass

    return fontsFound, lang_codes, lang_paragraphs


def getLangsInParagraphs(paragraphs, para_langs, lang_paragraphs):
    # Returns languages plus the text found for each with the log-probability
    langid.set_languages(['en', 'de','fr','it'])
    for p in paragraphs:
        if p.text == '':
            continue

        lang = langid.langid.classify(p.text)
        lang_code = lang[0]
        if lang_code in para_langs:
            para_langs[lang_code] += 1
            lang_paragraphs[lang_code].append([p.text, lang[1]])
        else:
            para_langs[lang_code] = 1
            lang_paragraphs[lang_code] = [p.text, lang[1]]
    return lang_paragraphs

def getFontsInParagraphs(paragraphs, fonts):
    for p in paragraphs:
        for r in p.runs:
            font = r.font.name
            if font in fonts:
                fonts[font] += 1
            else:
                fonts[font] = 1
    return fonts
            
#https://tedboy.github.io/flask/generated/flask.stream_with_context.html@
@app.route('/uploader/', methods = ['GET', 'POST'])
def upload_file():

    # True if conversion is requested, otherwise just doc info
    convertDoc = False
    get_doc_info = True

    lang = request.args.get('lang', 'und')

    who = '/uploader/%s' % lang
    
    if request.method: # anything should work!  == 'POST':
        formData = request.form.to_dict()
        if 'ConvertToUnicode' in formData:
            logger.debug('ConvertToUnicode')
            convertDoc = True
            get_doc_info = False
        logger.debug('DEBUG: %s', request.method)
        try:
            taskId = int(formData['taskId'])
        except:
            taskId = 117

        # Other options for conversion:
        remove_returns_in_block = False
        if 'remove_returns' in formData:
            remove_returns_in_block = formData['remove_returns']

        download_as_zip = False
        if 'folder_output' in formData:
            download_as_zip = formData['folder_output']

        use_vs = False
        if 'use_variation_selectors' in formData:
            use_vs = formData['use_variation_selectors']
        logger.debug('USE_VS = %s', use_vs)

        logger.debug('*** taskId = %d', taskId)
        # Get the information on the language 
        try:
            lang = formData['lang']
            lang_code = lang
            logger.debug('lang = %s', formData['lang'])
        except:
            lang_code = 'ff'
            lang = 'ff'
        logger.debug('FORMDATA = %s', formData)

        if 'file_path' in formData:
            inputFileName = formData['file_path']
        else:
            file = request.files['file']  # Filestorage object
            logger.debug('FILE = %s', file)
            inputFileName = file.filename

        logger.debug('inputFileName = %s' % inputFileName)
        if not inputFileName:
            return render_template('nofileselected.html', who=who)
            
        split_name = os.path.splitext(inputFileName)
        baseName = split_name[0]
        extension = split_name[1]
        outFileName = baseName + '_Unicode' + extension

        if extension == '.xlsx':
            cell_ranges = request.args.get('spreadsheet_region', None)
            converter = converters[lang]
            out_file_name = converter.get_outfile_name(inputFileName)  # Temporary
            debug_output = True
            processor = convertXls.convertWorkbook(inputFileName, out_file_name, converter, debug_output, cell_ranges=cell_ranges)

            processor.process()  # Do the requested conversion

            processor.workbook.save(out_file_name)
        elif extension == '.docx':
            # New thread for this id
            this_thread = exporting_threads[taskId] = ExportingThread()
            this_thread.start()
            this_thread.status = 'Creating doc %s from upload' % inputFileName

            doc, fileSize = createDocFromFile(file)

            if not doc:
                return render_template(
                    'error.html',
                    who=who,
                    error='%s %s' % ('Problem creating file', inputFileName))

            this_thread.status = 'Doc ready to process'

            fontsFound, para_langs, lang_paragraphs = findDocFonts(doc)
            if not convertDoc:
                # Just show information.
                return render_template(
                    'docinfo.html',
                    size="{:,}".format(fileSize),
                    filename=inputFileName,
                    paragraphs="{:,}".format(len(doc.paragraphs)),
                    sections=len(doc.sections),
                    tables=len(doc.tables),
                    fontDict=fontsFound,
                    para_langs=json.dumps(para_langs),
                    unicodeFont=formData['UnicodeFont']
                )

            this_thread.status = ('Paragraphs found: %d' % len(doc.paragraphs))


            # Call conversions on the document.
            langConverter = None
            logger.debug('LANG = %s', lang_code)
            try:
                langConverter = converters[lang_code]
            except KeyError:
                langConverter = None
                return render_template('unsupported_lang_code.html',
                                       lang_code=lang_code,
                                       supported_lang_codes= converts.keys()
                                       )

            langConverter.detectLang = langid.langid
            langConverter.ignoreLangs = ['en', 'fr']  # Not converted

            # Other settings
            langConverter.remove_returns_in_block = remove_returns_in_block
            langConverter.download_as_zip = download_as_zip
            langConverter.add_variant_selectors = use_vs

            langConverter.taskId = taskId

            newProgressObj = ProgressClass(langConverter, this_thread)

            try:
                try:
                    scriptIndex = int(formData['scriptIndex'])
                except:
                    logger.debug('NO SCRIPT INDEX')
                    scriptIndex = 0

                langConverter.setScriptIndex(scriptIndex)
                langConverter.setLowerMode(True)
                langConverter.setSentenceMode(True)
                paragraphs = doc.paragraphs
                print(' %s PARAGRAPHS' % len(paragraphs))
                msgToSend = '%d paragraphs in %s\n' % (len(paragraphs), inputFileName)
                countSent = 0
            except BaseException as err:
                return render_template('error.html',
                                       who=who,
                                       error='Bad langConverter: %s' % err)

            try:
                docConverter = ConvertDocx(langConverter, documentIn=doc,
                                           reportProgressObj=newProgressObj)
            except BaseException as error:
                print('Cannot create doc converter: %s' % error)
                return render_template('error.html', who=who, error=error)

            result = docConverter.processDocx()

            # This is the output .docx file
            # TODO: Show "saving"
            target_stream = BytesIO()
            result = doc.save(target_stream)

            # Download resulting converted document
            # Reset the pointer to the beginning.
            target_stream.seek(0)

            # Deal with non-ASCII in the file name
            # outFileName = fixNonAsciiFilename(outFileName)
            headerFileName = "attachment;filename=%s" % outFileName

            if download_as_zip:
                # Create zip file of .docx, word list, and info files.
                # Word list file
                wordFrequencies = langConverter.getSortedWordList()
                # Try to make this with a zip archive
                # Create a .tsv file of the word frequencies
                text_stream = StringIO()
                if wordFrequencies:
                    text_stream.write('%s\t%s\n' % ('Word', 'Times in file'))
                    for item in wordFrequencies:
                        outline = '%s\t%s\n' % (item[0], item[1])
                        text_stream.write(outline)
                text_stream.seek(0)
                wordsFileName = baseName + "_words.tsv"

                # Create an info file
                info_stream = StringIO()
                now = datetime.datetime.now()
                info_stream.write('Source filename = %s\n' % inputFileName)
                info_stream.write('Output filename = %s\n' % outFileName)
                info_stream.write('Converted to Unicode at %s\n' %
                                  now.strftime('%Y-%m-%d %H:%M:%S'))
                info_stream.write('File size:  {:,} bytes\n'.format(fileSize))
                info_stream.write('{:,} paragraphs\n'.format(len(doc.paragraphs)))
                info_stream.write(' %d sections\n' % len(doc.sections))
                info_stream.write(' %d tables\n' % len(doc.tables))
                info_stream.write(' fonts found = %s\n' % fontsFound)
                info_stream.write(' unicodeFont = %s\n' % formData['UnicodeFont'])
                info_stream.seek(0)

                # The zipfile contents
                zipStream = BytesIO()
                with ZipFile(zipStream, 'w') as zf:
                    zf.writestr(outFileName, target_stream.read())
                    zf.writestr(wordsFileName, text_stream.read())
                    zf.writestr('%s_info.txt' % baseName, info_stream.read())

                zipStream.seek(0)
                zipName = baseName + '_Unicode.zip'
                result_download = send_file(zipStream,
                                            as_attachment=True,
                                   download_name=zipName)
            else:
                # Just save the Unicode .docx file
                result_download = send_file(target_stream,
                                            as_attachment=True,
                                            download_name = outFileName)

            return result_download
        else:
            logger.error('!!! Not processing file %s !', inputFileName)
            return None

        # Done processing.
        newProgressObj.stop_updates('HALTING')


def createZipArchive(target_stream, headerFileName, baseName, wordFrequencies):
    # Try zip file...
    zipStream = BytesIO()
    zf= ZipFile(zipStream, 'w')
    try:
        zf.writestr(headerFileName, target_stream.read())
    except BaseException as err:
        print('*** Cannot put doc into zip file %s' % (err))
        return False

    text_stream = StringIO()
    for item in wordFrequencies:
        addLine = '%s\t%s\n' % (item[0], item[1])
    text_stream.write(addLine)
    text_stream.seek(0)
    frequenciesName = baseName + '_words.tsv'
    zf.writestr(frequenciesName, text_stream.read())
        
    return zipStream
        
@app.route('/testzip')
def testZip():
    document = Document()
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    target_stream = BytesIO()
    result = document.save(target_stream)
    target_stream.seek(0)
    headers = {
        "Content-Disposition":  'test_doc.docx'
    }

    # Create a text file
    text_stream = StringIO()
    text_stream.write("text for testing")
    text_stream.seek(0)
    
    zipStream = BytesIO()
    
    with ZipFile(zipStream, 'w') as zf:
        zf.writestr('testDoc.docx', target_stream.read())
        zf.writestr('textSample.txt', text_stream.read())

        zipHeaders = {
            "Content-Disposition":  'test_unicode.zip'
        }
    zipStream.seek(0)
    return send_file(zipStream, as_attachment=True,
                     download_name='testdoc.zip')

        # return Response(
        #     zipStream,
        #     # stream_with_context(read_file_chunks(zf)),
        #     mimetype="application/zip",
        #     headers=zipHeaders
        # )
    return 

def fixNonAsciiFilename(filename):
    charList = []
    for i in [ord(x) for x in filename]:
        if i < 128:
            charList.append(chr(i))
        else:
            charList.append('\\u%04x' % i)
    return ''.join(charList)

# get uploaded file into document form
def createDocFromFile(file):
    try:
        text = file.stream.read()
        data = BytesIO(text)
        count = len(text)
        doc = Document(data)
        data.close()
        return doc, count
    except BaseException as err:
        print('Cannot create Docx for %s. Err = %s' % (file, err))
        return None, -1

@app.route('/convertAdlam', methods = ['GET', 'POST'])
def convertAdlam():
    if request.method == 'POST':
        formData = request.form.to_dict()

        file = request.files['file']  # FileStorage object
        fileName = file.filename
        outFileName = os.path.splitext(fileName)[0] + '_Unicode.docx'

        doc, count = createDocFromFile(file)

        try:
            langConverter = adlamConversion.AdlamConverter()      

            #
            langConverter.setScriptIndex(0)
            langConverter.setLowerMode(True)
            langConverter.setSentenceMode(True)
            paragraphs = doc.paragraphs
            count = len(paragraphs)
            msgToSend = '%d paragraphs in %s\n' % (count, fileName)
            countSent = 0

        except BaseException as err:
            return 'Bad Adlam converter. Err = %s' % err
        
        try:
            docConverter = ConvertDocx(langConverter, documentIn=doc,
                                       reportProgressObj=newProgressObj)

            if docConverter:
                result = docConverter.processDocx()

                target_stream = BytesIO()
                result = doc.save(target_stream)          

                try:
                    wordFrequencies = converter.getSortedWordList()
                    if wordFrequencies:
                        # Do something with this information
                        words = convertwordFrequencies.keys()
                        for item in words:
                            print(word)
                except:
                    print('FAILED TO GET WORD LIST')
                    words = None
                
                # Download resulting converted document
                # Reset the pointer to the beginning.
                target_stream.seek(0)

                # Deal with non-ASCII in the file name
                outFileName = outFileName.encode(
                    'ascii', errors='backslashreplace')
                        
                headerFileName = "attachment;filename=%s" % outFileName

                headers = {
                    "Content-Disposition": headerFileName
                }
                try:
                    msgToSend += 'Starting download\n'
                    countSent = 1
                    return Response(
                        stream_with_context(read_file_chunks(target_stream)),
                        mimetype="application/vnd.openxmlformats-officnedocument.wordprocessingml.document",
                        headers=headers
                    )
                except BaseException as err:
                    print('**** Response download failure. Err = %s' % err)
                
        except BaseeException as err:
            return 'Conversion failed. with err %s' % err
        
#    return '%s file uploaded successfully with %d paragraphs' % (file.filename, count)

# https://stackoverflow.com/questions/12232304/how-to-implement-server-push-in-flask-framework
def event_stream():
    global msgToSend
    global countSent
    print('STREAM count = %d, msg=%s' % (countSent, msgToSend))
    if countSent < 1:
        print('STREAM msg:%s' % msgToSend)
        yield "data: {}\n\n".format(msgToSend);
        msgToSend = msgToSend + '.'
        countSent += 1
        
# One way to push data to client. Not working right.
# @app.route('/stream')
# def stream():
#     return Response(event_stream(), mimetype="text/event-stream")


# Set up a thread to allow status updates.
# https://codehunter.cc/a/flask/flask-app-update-progress-bar-while-function-runs
class ExportingThread(threading.Thread):
    def __init__(self):
        self.progress = 0
        super().__init__()
        self.status = "moving on"

    def run(self):
        # Your exporting stuff goes here ...
        while True:  # Wait for something in the queue.
            message = queue.get(block=True)
            self.progress += 10
            #print('THREAD QUEUE MESSAGE = %s' % message)
            #print('THREAD STATUS: %s' % (self.status))

    def setStatus(self, newMessage):
        self.status = newMessage

        
@app.route('/start')
def index():
    global exporting_threads

    thread_id = random.randint(0, 10000)
    exporting_threads[thread_id] = ExportingThread()
    exporting_threads[thread_id].start()

    print('THREAD ID = %d' % thread_id)
    return 'task id: #%s' % thread_id

@app.route('/progress/<int:thread_id>')
def progress(thread_id):
    global exporting_threads
    # Gets the latest status of the thread
    if thread_id in exporting_threads:
        return str(exporting_threads[thread_id].status)
    else:
        return str('Thread %s not found' % thread_id)


@app.route('/lang')
def testLangId():
    args = request.args
    text = args['text']
    if text:
        result = langid.classify(text)
        return str(result)
    else:
        return str("No text")

@app.route('/example_task_handler', methods=['POST'])
def example_task_handler():
    """Log the request payload."""
    payload = request.get_data(as_text=True) or '(empty payload)'
    print('Received task with payload: {}'.format(payload))
    return 'Printed task payload: {}'.format(payload)
# [END cloud_tasks_appengine_quickstart]


if __name__ == '__main__':
    # This is used when running locally only. When deploying to Google App
    # Engine, a webserver process such as Gunicorn will serve the app. This
    # can be configured by adding an `entrypoint` to app.yaml.
    app.run(host='127.0.0.1', port=8080, debug=True, threaded=True)
# [END gae_python37_app]


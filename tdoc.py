#!/usr/bin/python3
# -*- coding: utf-8 -*-

# Test setting doc from stream

from io import StringIO 

from io import BytesIO

from docx import Document

try:
    with open('a1.docx', 'rb') as f:
        source_stream = BytesIO(f.read())
    print('Source stream = %s' % source_stream)
    doc = Document(source_stream)

    print('Doc opens')
except BaseException as err:
    print("Document fails %s" % (err))

print('%d paragraphs' % len(doc.paragraphs))    
run = doc.add_paragraph().add_run()
run.text = "TEST TEXT"

print('%d paragraphs' % len(doc.paragraphs))    


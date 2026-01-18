#!/usr/bin/python3
#-*- coding: utf-8 -*-

# Read Unicode documents and check the data on the complex scripts found.

import glob
from io import BytesIO
import os
import sys

import docx
from docx import Document
from docx.oxml.shared import OxmlElement, qn

from convertDoc2 import ConvertDocx

# CONSTANT
default_langCode = 'phk'

# get uploaded file into document form
def createDocFromFile(file_path):
    try:
        file = open(file_path, 'rb')
        text = file.read()
        data = BytesIO(text)
        count = len(text)
        doc = Document(data)
        file.close()
        return doc, count
    except BaseException as err:
        print('Cannot create Docx for %s. Err = %s' % (file_path, err))
        return None, -1


def fix_cs_formatting_run(run_to_fix, user_cs_font_size, user_cs_font_name,
                          user_is_bold=None, langCode='phk', is_bidi=False):
    # Start solving the font size and name problem
    # https://stackoverflow.com/questions/45627652/python-docx-add-style-with-ctl-complex-text-layout-language
    #cs: complex script, ex, arabic

    rpr = run_to_fix.element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rpr.get_or_add_sz()
    szCs = OxmlElement('w:szCs')  # size
    sz= OxmlElement('w:sz')  # size
    rtl = OxmlElement('w:rtl')  # If this is a right-to-left
    rpr.append(szCs)
    rpr.append(sz)
    rpr.append(rtl)
    lang = OxmlElement('w:lang')  #language
    rpr.append(lang)
    if user_is_bold:
        # If bolding is desired
        bCs = OxmlElement('w:bCs')  #bold the complex language
        rpr.append(bCs)
        bCs.set(qn('w:val'), "True")
        b = OxmlElement('w:b')  # bold the english
        rpr.append(b)
        b.set(qn('w:val'), "True")
    sz.set(qn('w:val'), str(int(user_cs_font_size * 2)))
    szCs.set(qn('w:val'), str(int(user_cs_font_size * 2)))

    lang.set(qn("w:val"), langCode)
    lang.set(qn('w:bidi'), langCode)  # This depends on the language code

    rFonts.set(qn('w:cs'), user_cs_font_name)
    rFonts.set(qn('w:ascii'), user_cs_font_name)
    rFonts.set(qn('w:hAnsi'), user_cs_font_name)
    rFonts.set(qn('w:eastAsia'), user_cs_font_name)

def fix_paragraph_runs(para, user_cs_font_name=None, user_cs_font_size=12):
    runs = para.runs
    user_is_bold = False

    for run in runs:
        font_name = run.font.name
        if font_name == user_cs_font_name:
            fix_cs_formatting_run(run, user_cs_font_size, user_cs_font_name,
                                   user_is_bold)  # cs: complex script, ex, arabic

        continue

def checkComplex(lang, input_path, document=None, save_doc=False):
    if not Document:
        document, count = createDocFromFile(input_path)
    user_cs_font_size = 12
    user_cs_font_name = 'Arial'  # Default

    if lang == 'phk':
        user_cs_font_size = 12
        user_cs_font_name = 'PhakeRamayanaUnicode'

    # get the paragraphs
    paragraphs = document.paragraphs
    paragraphId = 0
    for para in paragraphs:
        # print('!!! Paragraph # %s: %s' % (paragraphId, para.text))
        # Now this is specialized for phk
        fixParagraphRuns(para, user_cs_font_name=user_cs_font_name, user_cs_font_size=user_cs_font_size)
        paragraphId += 1

    # Now check tables
    tables = document.tables
    tableId = 0
    for table in tables:
      #self.converter.current_table = table  # To help with setting font sizes.
      # print('!!! TABLE %d' % (tableId))
      tableId += 1
      # if self.progressObj:
      #   self.progressObj.send('Table %d, %d rows' % (tableId, len(table.rows)))
      row_id = 0
      rows = table.rows
      for row in rows:
        row_id += 1
        for cell in row.cells:
          paragraphs = cell.paragraphs
          for para in paragraphs:
              fixParagraphRuns(para)

    if save_doc:
        new_doc_name  = input_path.replace('_Unicode', '_UnicodeFixed')
        document.save(new_doc_name)

    return

def main(argv):
    if len(argv) < 3:
        print('Convert .docx files from font encodings to Unicode text')
        print('Usage: python3 command_line lang_code file1 file2 file ...')
        return

    lang = argv[1]

    # For each item in the list, [2:...]
    files = []
    for doc_path in argv[2:]:
        if os.path.isdir(doc_path):
            # Expand with glob
            files.extend(glob.glob(doc_path + "/*.docx"))
        else:
            files.append(doc_path)

    for file_path in files:
        unicode_in_name = file_path.find('_Unicode.')
        if unicode_in_name < 0:
            # Only look at Unicode converted files
            continue
        print('Checking complex scripts %s in document %s' % (lang, file_path))
        checkComplex(lang, file_path, save_doc=True)


if __name__ == '__main__':
    main(sys.argv)
